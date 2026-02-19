# app/services/loader_service.py
from typing import List, Tuple
from pathlib import Path
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
import chardet


class LoaderService:
    """문서 로더 서비스"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".py", ".js"}

    def load(self, file_path: str) -> Tuple[str, str]:
        """파일 로드 (내용, 타입)"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"지원하지 않는 형식: {ext}")

        if ext == ".pdf":
            return self._load_pdf(path), "pdf"
        elif ext == ".docx":
            return self._load_docx(path), "docx"
        else:
            return self._load_text(path), ext[1:]

    def load_bytes(self, content: bytes, filename: str) -> Tuple[str, str]:
        """바이트에서 로드"""
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._load_pdf_bytes(content), "pdf"
        elif ext == ".docx":
            return self._load_docx_bytes(content), "docx"
        else:
            return self._load_text_bytes(content), ext[1:] if ext else "txt"

    def _load_text(self, path: Path) -> str:
        """텍스트 파일 로드"""
        with open(path, "rb") as f:
            raw = f.read()

        # 인코딩 감지
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8")

        return raw.decode(encoding)

    def _load_text_bytes(self, content: bytes) -> str:
        """바이트에서 텍스트 로드"""
        detected = chardet.detect(content)
        encoding = detected.get("encoding", "utf-8")
        return content.decode(encoding)

    def _load_pdf(self, path: Path) -> str:
        """PDF 로드"""
        reader = PdfReader(path)
        texts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

        return "\n\n".join(texts)

    def _load_pdf_bytes(self, content: bytes) -> str:
        """바이트에서 PDF 로드"""
        import io
        reader = PdfReader(io.BytesIO(content))
        texts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

        return "\n\n".join(texts)

    def _load_docx(self, path: Path) -> str:
        """DOCX 로드"""
        doc = DocxDocument(path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text)

    def _load_docx_bytes(self, content: bytes) -> str:
        """바이트에서 DOCX 로드"""
        import io
        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text)
